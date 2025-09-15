import csv
import mysql.connector 
from docx import Document
from docx.shared  import Cm, Pt, Inches
from docx.enum.text  import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns  import qn
from bs4 import BeautifulSoup 
from docx.oxml  import OxmlElement
from datetime import datetime 
 
def get_db_connection():
    """建立数据库连接"""
    try:
        conn = mysql.connector.connect( 
            host='localhost',
            user='root',
            password='root',
            database='teaching_assistant'
        )
        return conn 
    except mysql.connector.Error  as err:
        print(f"数据库连接错误: {err}")
        return None
 
def get_student_answers(student_name, target_date=None):
    """获取指定学生在特定日期的答案"""
    conn = get_db_connection()
    if not conn:
        return None
    
    try:
        cursor = conn.cursor(dictionary=True) 
        
        # 查询学生ID 
        cursor.execute("SELECT  student_id FROM students WHERE student_name = %s", (student_name,))
        student = cursor.fetchone() 
        if not student:
            print(f"未找到学生: {student_name}")
            return None
        
        # 构建查询语句 
        query = """
        SELECT q.answer_text,  sr.performance_score,  sr.student_feedback  
        FROM student_records sr
        JOIN questions q ON sr.question_id  = q.question_id  
        WHERE sr.student_id  = %s 
        """
        params = [student['student_id']]
        
        # 添加日期筛选条件
        if target_date:
            try:
                # 将输入日期转换为数据库格式
                date_obj = datetime.strptime(target_date,  "%Y/%m/%d")
                query += " AND DATE(sr.create_at)  = %s"
                params.append(date_obj.strftime("%Y-%m-%d")) 
            except ValueError:
                print(f"日期格式错误，使用默认格式: {target_date}")
                return None 
        
        query += " ORDER BY sr.create_at" 
        
        cursor.execute(query,  params)
        return cursor.fetchall() 
        
    except mysql.connector.Error  as err:
        print(f"数据库查询错误: {err}")
        return None 
    finally:
        conn.close() 
 
def add_sup_sub(run, tag):
    """添加上标或下标"""
    r = run._r 
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'),  'superscript' if tag.name  == 'sup' else 'subscript')
    rPr = r.get_or_add_rPr() 
    rPr.append(vertAlign) 
    
    original_size = run.font.size  or Pt(9)
    new_size = max(original_size - Pt(2), Pt(7))
    run.font.size  = new_size
    
    position = OxmlElement('w:position')
    position.set(qn('w:val'),  "6" if tag.name  == 'sup' else "-4")
    rPr.append(position) 
 
def process_html_answer(doc, html_content):
    """处理HTML格式的答案并添加到文档"""
    if not html_content:
        return 
    
    soup = BeautifulSoup(html_content, 'html.parser') 
    
    # 处理有序列表 
    for ol in soup.find_all('ol'): 
        list_num = 0
        for li in ol.find_all('li',  recursive=False):
            list_num += 1
            li_p = doc.add_paragraph() 
            li_run = li_p.add_run(f"({list_num})  ")
            li_run.font.name  = 'Times New Roman'  # 确保数字和字母使用英文字体 
            for content in li.contents: 
                if content.name  in ['sup', 'sub']:
                    run = li_p.add_run(content.text) 
                    run.font.name  = 'Times New Roman'  # 确保上标/下标使用英文字体
                    add_sup_sub(run, content)
                elif content.name  == 'i':  # 处理斜体标签
                    run = li_p.add_run(content.text) 
                    run.font.italic  = True
                    run.font.name  = 'Times New Roman'
                elif content.name  is None:
                    text = str(content)
                    # 替换特殊字符
                    replacements = {
                        '×': '×', '÷': '÷', '°': '°', 
                        '·': '·', '−': '-', '≈': '≈'
                    }
                    for k, v in replacements.items(): 
                        text = text.replace(k,  v)
                    run = li_p.add_run(text) 
                    run.font.name  = 'Times New Roman'  # 确保字母使用英文字体 
    
    # 处理段落 
    for para in soup.find_all('p'): 
        if not para.text.strip()  or para.find_parent('ol'): 
            continue
            
        p = doc.add_paragraph() 
        for content in para.contents: 
            if content.name  in ['sup', 'sub']:
                run = p.add_run(content.text) 
                run.font.name  = 'Times New Roman'  # 确保上标/下标使用英文字体
                add_sup_sub(run, content)
            elif content.name  == 'i':  # 处理斜体标签
                run = p.add_run(content.text) 
                run.font.italic  = True 
                run.font.name  = 'Times New Roman'
            elif content.name  is None:
                text = str(content)
                replacements = {
                    '×': '×', '÷': '÷', '°': '°', 
                    '·': '·', '−': '-', '≈': '≈'
                }
                for k, v in replacements.items(): 
                    text = text.replace(k,  v)
                run = p.add_run(text) 
                run.font.name  = 'Times New Roman'  # 确保字母使用英文字体
 
def create_answer_document(students_data, target_date=None):
    """创建答案文档"""
    doc = Document()
    
    # 设置页面布局 - 左边距1cm，右边距7cm 
    section = doc.sections[0] 
    section.left_margin  = Cm(15)
    section.right_margin  = Cm(1)
    
    # 设置文档默认字体 
    style = doc.styles['Normal'] 
    font = style.font 
    font.name  = '宋体'
    font.size  = Pt(9)  # 小五
    
    # 添加每个学生的答案
    answer_counter = 0  # 用于跟踪当前页面的答案数量 
    
    for student_name, answers in students_data.items(): 
        if not answers:
            continue
        
        # 添加每个问题的答案
        for i, answer in enumerate(answers, 1):
            # 每两个答案后添加分页符（除了第一个答案）
            if answer_counter == 2:
                doc.add_page_break() 
                answer_counter = 0
            
            # 添加答案标题 
            answer_title = doc.add_paragraph() 
            answer_run = answer_title.add_run(f' 答案 {i} - {student_name}')
            answer_run.font.size  = Pt(9)
            answer_run.font.name  = '宋体'  # 明确设置中文字体
            
            # 处理HTML格式的答案 
            if answer['answer_text']:
                process_html_answer(doc, answer['answer_text'])
            
            # 添加评分和反馈（如果有）
            if answer['performance_score']:
                p_score = doc.add_paragraph(f' 评分: {answer["performance_score"]}')
                p_score.runs[0].font.name  = '宋体'  # 中文部分使用宋体 
            
            if answer['student_feedback']:
                p_feedback = doc.add_paragraph(f' 反馈: {answer["student_feedback"]}')
                p_feedback.runs[0].font.name  = '宋体'  # 中文部分使用宋体
            
            # 如果不是最后一个答案，添加分隔线
            if answer_counter < 1 and i < len(answers):
                doc.add_paragraph("_"*50).alignment  = WD_PARAGRAPH_ALIGNMENT.CENTER 
            
            answer_counter += 1
    
    # 保存文档
    filename = f"学生答案汇总_{target_date.replace('/',  '-')}.docx" if target_date else "学生答案汇总.docx"
    doc.save(filename) 
    print(f"已生成: {filename}")
 
def main():
    # 读取批改顺序CSV文件 
    try:
        with open('批改顺序.csv', mode='r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            students = [row['姓名'] for row in reader]
    except FileNotFoundError:
        print("错误: 找不到'批改顺序.csv'文件")
        return
    except Exception as e:
        print(f"读取CSV文件错误: {e}")
        return
    
    # 获取用户输入的日期
    target_date = input("请输入要筛选的日期(格式: YYYY/MM/DD，直接回车表示不筛选): ").strip()
    if target_date:
        try:
            # 验证日期格式 
            datetime.strptime(target_date,  "%Y/%m/%d")
        except ValueError:
            print("错误: 日期格式不正确，请使用YYYY/MM/DD格式")
            return
    
    # 收集所有学生的答案数据 
    students_data = {}
    for student in students:
        print(f"正在处理: {student}")
        answers = get_student_answers(student, target_date)
        students_data[student] = answers if answers else None
    
    # 生成答案文档 
    create_answer_document(students_data, target_date if target_date else None)
 
if __name__ == "__main__":
    main()