#优化：优先选择学生没有做过的题目出题，如果都做过，则随机出题。
from flask import Flask, request, jsonify, send_from_directory
import mysql.connector
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import os
from docx.oxml import OxmlElement
import random
from datetime import datetime

# 创建Flask应用，不指定static_folder
app = Flask(__name__)

# 设置数据库连接信息（根据实际情况修改）
DB_CONFIG = {
    'host': 'localhost',
    'user': 'root',
    'password': 'root',
    'database': 'teaching_assistant'
}

# 添加CORS处理
@app.after_request
def add_cors_headers(response):
    # 允许HTTPS来源
    allowed_origins = ['https://localhost', 'https://yourdomain.com']
    origin = request.headers.get('Origin')
    if origin in allowed_origins:
        response.headers['Access-Control-Allow-Origin'] = origin
    
    # 其他必要的CORS头
    if request.method == 'OPTIONS':
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        response.headers['Access-Control-Allow-Credentials'] = 'true'
    return response

# 处理OPTIONS预检请求
@app.route('/get-students', methods=['OPTIONS'])
def options_handler():
    return '', 200

@app.route('/get-students')
def get_students():
    """获取学生列表"""
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT student_id, student_name FROM students ORDER BY student_name")
        students = cursor.fetchall()
        
        return jsonify(students)
    except Exception as e:
        print(f"获取学生数据错误: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if 'conn' in locals() and conn.is_connected():
            cursor.close()
            conn.close()

def get_student_id_by_name(student_name):
    """根据学生姓名获取学生ID"""
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT student_id FROM students WHERE student_name = %s", (student_name,))
        result = cursor.fetchone()
        
        return result['student_id'] if result else None
    except Exception as e:
        print(f"获取学生ID错误: {e}")
        return None
    finally:
        if 'conn' in locals() and conn.is_connected():
            cursor.close()
            conn.close()

def add_sup_sub(run, tag):
    """完美添加上标或下标"""
    r = run._r
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript' if tag.name == 'sup' else 'subscript')
    rPr = r.get_or_add_rPr()
    rPr.append(vertAlign)
    
    original_size = run.font.size or Pt(12)
    new_size = max(original_size - Pt(2), Pt(8))
    run.font.size = new_size
    
    position = OxmlElement('w:position')
    position.set(qn('w:val'), "6" if tag.name == 'sup' else "-4")
    rPr.append(position)

def add_question_image(doc, image_path):
    """添加题目图片到文档"""
    if not image_path:
        return
        
    abs_path = os.path.abspath(image_path)
    if not os.path.exists(abs_path):
        abs_path = os.path.abspath(os.path.join(os.path.dirname(__file__), image_path))
    
    if os.path.exists(abs_path):
        try:
            doc.add_picture(abs_path, height=Inches(1.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph()
        except Exception as e:
            print(f"无法添加图片 {abs_path}: {e}")
    else:
        print(f"图片不存在: {abs_path}")

def process_html_tables(soup, doc):
    """处理HTML中的表格并添加到Word文档"""
    for table in soup.find_all('table'):
        # 先确定行数和列数
        rows = len(table.find_all('tr'))
        cols = max(len(tr.find_all(['th', 'td'])) for tr in table.find_all('tr')) if rows > 0 else 1
        
        # 创建Word表格
        word_table = doc.add_table(rows=rows, cols=cols)
        word_table.style = 'Table Grid'
        
        # 添加表格内容
        for row_idx, tr in enumerate(table.find_all('tr')):
            cells = tr.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                if row_idx < rows and col_idx < cols:  # 确保不越界
                    word_table.cell(row_idx, col_idx).text = cell.get_text()
        
        doc.add_paragraph()  # 表格后添加空行

def get_question_from_db(topic, difficulty_level, student_id=None, exclude_ids=None):
    """从数据库获取题目（严格遵循优先未做过逻辑）
    参数:
        topic: 题目主题 
        difficulty_level: 难度级别 
        student_id: 学生ID（可选）
        exclude_ids: 要排除的题目ID列表（可选）
    返回:
        随机选择的题目字典 或 None（无符合条件题目时）
    """
    try:
        conn = mysql.connector.connect(**DB_CONFIG) 
        cursor = conn.cursor(dictionary=True) 
        
        # 基础查询条件（始终生效）
        base_conditions = [
            "q.topic  = %s",
            "q.difficulty_level  = %s"
        ]
        params = [topic, difficulty_level]
        
        # 添加排除ID条件 
        if exclude_ids:
            base_conditions.append(f"q.question_id  NOT IN ({','.join(['%s']*len(exclude_ids))})")
            params.extend(exclude_ids) 
        
        # 阶段1：优先查询学生未做过的题目（如果提供了student_id）
        if student_id:
            # 构造完整查询语句 
            query = f"""
                SELECT q.question_id,  q.question_text,  q.question_image  
                FROM questions q 
                LEFT JOIN student_records sr ON (
                    q.question_id  = sr.question_id  
                    AND sr.student_id  = %s 
                )
                WHERE {' AND '.join(base_conditions)}
                AND sr.question_id  IS NULL  # 确保学生没做过 
                ORDER BY RAND()
                LIMIT 1 
            """
            cursor.execute(query,  [student_id] + params)
            unused_question = cursor.fetchone() 
            
            if unused_question:
                return unused_question 
        
        # 阶段2：普通查询（学生未提供或所有未做过题目已用完）
        query = f"""
            SELECT q.question_id,  q.question_text,  q.question_image  
            FROM questions q 
            WHERE {' AND '.join(base_conditions)}
            ORDER BY RAND()
            LIMIT 1 
        """
        cursor.execute(query,  params)
        return cursor.fetchone() 
        
    except Exception as e:
        print(f"数据库查询错误: {e}")
        return None 
    finally:
        if 'conn' in locals() and conn.is_connected(): 
            cursor.close() 
            conn.close() 

def add_header_with_timestamp(doc):
    """添加带有时间戳的页眉"""
    # 设置页面边距：左边距1cm，右边距7cm
    section = doc.sections[0]
    section.left_margin = Inches(0.3937)  # 1cm ≈ 0.3937英寸
    section.right_margin = Inches(2.7559)  # 7cm ≈ 2.7559英寸
    
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # 添加生成时间
    run = paragraph.add_run(f"试卷生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 修改 add_student_questions 函数，使其处理任意数量的题目
def add_student_questions(doc, student_name, selections, question_ids, is_first_student=False):
    """为单个学生添加试卷内容（完整修改版）
    参数:
        doc: Word文档对象
        student_name: 学生姓名
        selections: 题目设置列表 [{"topic":主题, "difficulty":难度}, ...]
        question_ids: 存储题目ID的字典 {学生姓名: [题目信息]}
        is_first_student: 是否是第一个学生（用于控制分页）
    """
    # 获取学生ID
    student_id = get_student_id_by_name(student_name)
    if not student_id:
        print(f"无法找到学生ID: {student_name}")
        return question_ids
    
    # 如果不是第一个学生，添加分页符（分隔不同学生的试卷）
    if not is_first_student:
        # 检查文档最后一页是否有内容
        if len(doc.paragraphs) > 0:
            last_paragraph = doc.paragraphs[-1]
            if any(run.text.strip() for run in last_paragraph.runs):
                doc.add_page_break() 
    
    # ========== 第一部分：学生标题 ==========
    # 添加学生姓名标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f"{student_name}")
    title_run.bold = True
    title_run.font.size = Pt(20)
    
    
    # 如果题目数量超过2道，则每页放2道题
    if len(selections) > 2:
        # 计算需要多少页
        num_pages = (len(selections) + 1) // 2
        
        for page in range(num_pages):
            # 如果不是第一页，添加分页符
            if page > 0:
                doc.add_page_break()

            # 当前页的两道题
            start_idx = page * 2
            end_idx = min(start_idx + 2, len(selections))
            
            for i in range(start_idx, end_idx):
                # 获取已选题目ID用于去重
                exclude_ids = [q['question_id'] for q in question_ids[student_name]]
                
                # 获取题目
                question = get_question_from_db(
                    selections[i]['topic'],
                    selections[i]['difficulty'],
                    student_id,
                    exclude_ids
                )
                
                if question:
                    # 添加题号
                    p = doc.add_paragraph()
                    run = p.add_run(f"{i+1}、")
                    run.bold = True
                    
                    # 处理题目文本（HTML格式）
                    if question['question_text']:
                        soup = BeautifulSoup(question['question_text'], 'html.parser')
                        
                        # 处理表格
                        process_html_tables(soup, doc)
                        
                        # 处理有序列表
                        for ol in soup.find_all('ol'):
                            list_num = 0
                            for li in ol.find_all('li', recursive=False):
                                list_num += 1
                                li_p = doc.add_paragraph()
                                li_run = li_p.add_run(f"（{list_num}） ")
                                for content in li.contents:
                                    if content.name in ['sup', 'sub']:
                                        run = li_p.add_run(content.text)
                                        add_sup_sub(run, content)
                                    elif content.name is None:
                                        text = str(content)
                                        # 替换特殊符号
                                        text = text.replace('×', '×').replace('÷', '÷')
                                        text = text.replace('°', '°').replace('−', '-')
                                        li_p.add_run(text)
                        
                        # 处理普通段落
                        for para in soup.find_all('p'):
                            if not para.text.strip() or para.find_parent('ol'):
                                continue
                                
                            for content in para.contents:
                                if content.name in ['sup', 'sub']:
                                    run = p.add_run(content.text)
                                    add_sup_sub(run, content)
                                elif content.name is None:
                                    text = str(content)
                                    text = text.replace('×', '×').replace('÷', '÷')
                                    text = text.replace('°', '°').replace('−', '-')
                                    p.add_run(text)
                    
                    # 添加题目图片
                    if question['question_image']:
                        add_question_image(doc, question['question_image'])
                    
                    # 记录题目信息
                    question_ids[student_name].append({
                        'question_num': i+1,
                        'question_id': question['question_id'],
                        'topic': selections[i]['topic'],
                        'difficulty': selections[i]['difficulty']
                    })
                    
                    # 两道题之间添加空行
                    if i < end_idx - 1:
                        doc.add_paragraph()
    else:
        # 题目数量≤2道，每道题单独一页
        for i in range(len(selections)):
            # 如果不是第一道题，添加分页符
            if i > 0:
                doc.add_page_break()
            
            # 获取已选题目ID用于去重
            exclude_ids = [q['question_id'] for q in question_ids[student_name]]
            
            # 获取题目
            question = get_question_from_db(
                selections[i]['topic'],
                selections[i]['difficulty'],
                student_id,
                exclude_ids
            )
            
            if question:
                # 添加题号
                p = doc.add_paragraph()
                run = p.add_run(f"{i+1}、")
                run.bold = True
                
                # 处理题目文本（与之前相同的逻辑）
                if question['question_text']:
                    soup = BeautifulSoup(question['question_text'], 'html.parser')
                    process_html_tables(soup, doc)
                    
                    for ol in soup.find_all('ol'):
                        list_num = 0
                        for li in ol.find_all('li', recursive=False):
                            list_num += 1
                            li_p = doc.add_paragraph()
                            li_run = li_p.add_run(f"（{list_num}） ")
                            for content in li.contents:
                                if content.name in ['sup', 'sub']:
                                    run = li_p.add_run(content.text)
                                    add_sup_sub(run, content)
                                elif content.name is None:
                                    text = str(content)
                                    text = text.replace('×', '×').replace('÷', '÷')
                                    text = text.replace('°', '°').replace('−', '-')
                                    li_p.add_run(text)
                    
                    for para in soup.find_all('p'):
                        if not para.text.strip() or para.find_parent('ol'):
                            continue
                            
                        for content in para.contents:
                            if content.name in ['sup', 'sub']:
                                run = p.add_run(content.text)
                                add_sup_sub(run, content)
                            elif content.name is None:
                                text = str(content)
                                text = text.replace('×', '×').replace('÷', '÷')
                                text = text.replace('°', '°').replace('−', '-')
                                p.add_run(text)
                
                # 添加题目图片
                if question['question_image']:
                    add_question_image(doc, question['question_image'])
                
                # 记录题目信息
                question_ids[student_name].append({
                    'question_num': i+1,
                    'question_id': question['question_id'],
                    'topic': selections[i]['topic'],
                    'difficulty': selections[i]['difficulty']
                })
    
    return question_ids


def save_to_student_records(student_name, question_ids):
    """将出卷记录保存到student_records表"""
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        student_id = get_student_id_by_name(student_name)
        if not student_id:
            print(f"无法找到学生ID: {student_name}")
            return False
        
        for question in question_ids:
            # 检查是否已存在相同记录
            cursor.execute("""
                SELECT record_id FROM student_records 
                WHERE student_id = %s AND question_id = %s
            """, (student_id, question['question_id']))
            
            if not cursor.fetchone():
                # 不存在则插入新记录
                cursor.execute("""
                    INSERT INTO student_records 
                    (student_id, question_id, create_at)
                    VALUES (%s, %s, %s)
                """, (student_id, question['question_id'], datetime.now()))
        
        conn.commit()
        return True
    except Exception as e:
        print(f"保存到student_records表失败: {e}")
        conn.rollback()
        return False
    finally:
        if 'conn' in locals() and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/')
def index():
    """返回前端页面"""
    # 直接从当前目录发送index.html文件
    return send_from_directory(os.path.dirname(__file__), 'index.html')

@app.route('/generate', methods=['POST'])
def generate_exam():
    """处理生成试卷请求"""
    try:
        data = request.json
        
        # 验证输入数据
        if not data.get('students') or len(data['students']) == 0:
            return jsonify({'success': False, 'error': '请至少选择一名学生'}), 400
        
        if not data.get('selections') or len(data['selections']) == 0:
            return jsonify({'success': False, 'error': '请至少设置一道题目'}), 400
        
        # 处理输出文件名
        output_file = data['output_file']
        if not output_file.endswith('.docx'):
            output_file += '.docx'
        
        desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
        output_path = os.path.abspath(os.path.join(desktop_path, output_file))
        
        # 检查文件是否存在
        file_exists = os.path.exists(output_path)
        
        # 创建或加载Word文档
        if file_exists:
            # 文件存在则加载现有文档
            doc = Document(output_path)
            
            # 确保文档有至少一个section
            if not doc.sections:
                doc.add_section()
            
            # 设置页面边距：左边距1cm，右边距7cm
            section = doc.sections[0]
            section.left_margin = Inches( 0.3937)
            section.right_margin = Inches(2.7559)
            
            # 检查文档是否以分页符结束
            needs_page_break = False
            if len(doc.paragraphs) > 0:
                last_paragraph = doc.paragraphs[-1]
                if any(run.text.strip() for run in last_paragraph.runs):
                    needs_page_break = True
            
            # 只在需要时添加分页符
            if needs_page_break:
                # 添加真正的分页符
                doc.add_page_break()
                # 移除可能自动添加的空段落
                if len(doc.paragraphs) > 1 and not doc.paragraphs[-1].runs:
                    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
        else:
            # 文件不存在则创建新文档
            doc = Document()
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加带有时间戳的页眉
            add_header_with_timestamp(doc)
        
        # 为每个学生生成试卷
        question_ids = {student: [] for student in data['students']}
        
        for i, student in enumerate(data['students']):
            # 修改2：简化分页逻辑，只在需要时添加分页符
            if i > 0 or (file_exists and i == 0 and needs_page_break):
                # 确保分页符是有效的
                doc.add_page_break()
                # 移除可能自动添加的空段落
                if len(doc.paragraphs) > 1 and not doc.paragraphs[-1].runs:
                    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            
            question_ids = add_student_questions(doc, student, data['selections'], question_ids, is_first_student=(i==0 and not file_exists))
        
        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        # 将出卷记录保存到数据库
        for student, questions in question_ids.items():
            if not save_to_student_records(student, questions):
                print(f"警告: 未能保存学生 {student} 的出卷记录")
        
        # 准备返回的question_ids信息
        formatted_question_ids = {}
        for student, questions in question_ids.items():
            formatted_question_ids[student] = [
                {
                    'question_num': q['question_num'],
                    'question_id': q['question_id'],
                    'topic': q['topic'],
                    'difficulty': q['difficulty']
                }
                for q in questions
            ]
        
        return jsonify({
            'success': True,
            'output_path': output_path,
            'question_ids': formatted_question_ids,
            'appended': file_exists  # 指示是否是追加操作
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

if __name__ == "__main__":
    app.run(debug=True, port=8080)